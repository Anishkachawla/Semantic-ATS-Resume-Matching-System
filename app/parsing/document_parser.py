"""
Document parsing: turns an uploaded file (PDF, DOCX, or .txt) into plain text.

This stage deliberately does nothing "smart" — it's pure text extraction.
Section detection happens downstream in section_extractor.py, kept separate
so you can swap either piece independently (e.g. a better PDF extractor
later without touching section logic).
"""
import os

import pdfplumber
from docx import Document


def extract_text_from_pdf(path: str) -> str:
    """Extract text from a PDF, page by page."""
    chunks = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text() or ""
            chunks.append(page_text)
    return "\n".join(chunks)


def extract_text_from_docx(path: str) -> str:
    """Extract text from a .docx, including table cells (resumes often use tables)."""
    doc = Document(path)
    chunks = [p.text for p in doc.paragraphs]

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    chunks.append(cell.text)

    return "\n".join(chunks)


def extract_text_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path: str) -> str:
    """Dispatch to the right extractor based on file extension."""
    ext = os.path.splitext(path)[1].lower()

    if ext == ".pdf":
        return extract_text_from_pdf(path)
    elif ext == ".docx":
        return extract_text_from_docx(path)
    elif ext == ".txt":
        return extract_text_from_txt(path)
    else:
        raise ValueError(
            f"Unsupported file type: '{ext}'. Supported: .pdf, .docx, .txt"
        )
